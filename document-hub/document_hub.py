#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Document Hub - 统一文档处理中心
提供Word、Excel、PDF、PPT的读取、生成和转换功能
"""

import os
import sys
from pathlib import Path
from typing import Optional, List, Dict, Any, Union
from dataclasses import dataclass
from enum import Enum

class DocumentType(Enum):
    WORD = "word"
    EXCEL = "excel"
    PDF = "pdf"
    PPT = "ppt"
    MARKDOWN = "markdown"
    AUDIO = "audio"
    VIDEO = "video"
    UNKNOWN = "unknown"

class DocumentError(Exception):
    """文档处理异常基类"""
    pass

class DocumentReadError(DocumentError):
    """文档读取异常"""
    pass

class DocumentWriteError(DocumentError):
    """文档写入异常"""
    pass

class DocumentConvertError(DocumentError):
    """文档转换异常"""
    pass

@dataclass
class DocumentInfo:
    """文档信息"""
    path: str
    doc_type: DocumentType
    size: int
    pages: Optional[int] = None
    paragraphs: Optional[int] = None
    tables: Optional[int] = None

def get_document_type(file_path: str) -> DocumentType:
    """根据文件扩展名判断文档类型"""
    ext = Path(file_path).suffix.lower()
    type_map = {
        '.docx': DocumentType.WORD,
        '.doc': DocumentType.WORD,
        '.xlsx': DocumentType.EXCEL,
        '.xls': DocumentType.EXCEL,
        '.pdf': DocumentType.PDF,
        '.pptx': DocumentType.PPT,
        '.ppt': DocumentType.PPT,
        '.md': DocumentType.MARKDOWN,
        '.markdown': DocumentType.MARKDOWN,
        # 音频格式
        '.mp3': DocumentType.AUDIO,
        '.wav': DocumentType.AUDIO,
        '.aac': DocumentType.AUDIO,
        '.flac': DocumentType.AUDIO,
        '.ogg': DocumentType.AUDIO,
        '.m4a': DocumentType.AUDIO,
        '.wma': DocumentType.AUDIO,
        # 视频格式
        '.mp4': DocumentType.VIDEO,
        '.avi': DocumentType.VIDEO,
        '.mov': DocumentType.VIDEO,
        '.mkv': DocumentType.VIDEO,
        '.flv': DocumentType.VIDEO,
        '.wmv': DocumentType.VIDEO,
        '.webm': DocumentType.VIDEO,
        '.m4v': DocumentType.VIDEO,
    }
    return type_map.get(ext, DocumentType.UNKNOWN)

def sanitize_path(path: str) -> str:
    """脱敏路径，移除敏感信息"""
    # 移除用户主目录路径
    home = str(Path.home())
    if path.startswith(home):
        path = path.replace(home, "~")
    # 移除workspace绝对路径
    workspace = "/Users/delta/.openclaw/workspace"
    if path.startswith(workspace):
        path = path.replace(workspace, "workspace")
    return path

class MediaHub:
    """音视频处理中心"""
    
    def __init__(self):
        self.supported_audio_formats = ['.mp3', '.wav', '.aac', '.flac', '.ogg', '.m4a', '.wma']
        self.supported_video_formats = ['.mp4', '.avi', '.mov', '.mkv', '.flv', '.wmv', '.webm', '.m4v']
    
    def convert_audio(self, input_path: str, output_path: str, **options) -> str:
        """
        转换音频格式
        
        Args:
            input_path: 输入音频文件路径
            output_path: 输出音频文件路径
            **options: 转换选项
                - bitrate: 比特率 (如 "128k", "192k", "320k")
                - sample_rate: 采样率 (如 44100, 48000)
                - channels: 声道数 (1=单声道, 2=立体声)
        
        Returns:
            输出文件的绝对路径
        """
        try:
            from pydub import AudioSegment
            
            # 加载音频
            audio = AudioSegment.from_file(input_path)
            
            # 应用选项
            if 'sample_rate' in options:
                audio = audio.set_frame_rate(options['sample_rate'])
            if 'channels' in options:
                audio = audio.set_channels(options['channels'])
            
            # 导出
            export_kwargs = {}
            if 'bitrate' in options:
                export_kwargs['bitrate'] = options['bitrate']
            
            audio.export(output_path, **export_kwargs)
            return os.path.abspath(output_path)
            
        except ImportError:
            raise DocumentConvertError("pydub未安装，无法转换音频。请运行: pip install pydub")
        except Exception as e:
            raise DocumentConvertError(f"音频转换失败: {str(e)}")
    
    def convert_video(self, input_path: str, output_path: str, **options) -> str:
        """
        转换视频格式
        
        Args:
            input_path: 输入视频文件路径
            output_path: 输出视频文件路径
            **options: 转换选项
                - fps: 帧率
                - resolution: 分辨率 (如 "1920x1080", "1280x720")
                - bitrate: 视频比特率
                - audio: 是否保留音频 (True/False)
                - codec: 视频编码器 (如 "libx264", "libx265")
        
        Returns:
            输出文件的绝对路径
        """
        try:
            from moviepy.editor import VideoFileClip
            
            # 加载视频
            video = VideoFileClip(input_path)
            
            # 应用选项
            if 'fps' in options:
                video = video.set_fps(options['fps'])
            if 'resolution' in options:
                width, height = map(int, options['resolution'].split('x'))
                video = video.resize(newsize=(width, height))
            if options.get('audio') is False:
                video = video.without_audio()
            
            # 导出
            export_kwargs = {'logger': None}  # 禁用日志
            if 'codec' in options:
                export_kwargs['codec'] = options['codec']
            if 'bitrate' in options:
                export_kwargs['bitrate'] = options['bitrate']
            
            video.write_videofile(output_path, **export_kwargs)
            video.close()
            
            return os.path.abspath(output_path)
            
        except ImportError:
            raise DocumentConvertError("moviepy未安装，无法转换视频。请运行: pip install moviepy")
        except Exception as e:
            raise DocumentConvertError(f"视频转换失败: {str(e)}")
    
    def extract_audio_from_video(self, video_path: str, output_path: str, **options) -> str:
        """
        从视频提取音频
        
        Args:
            video_path: 视频文件路径
            output_path: 输出音频文件路径
            **options: 提取选项
                - format: 音频格式 (如 "mp3", "wav", "aac")
                - bitrate: 比特率
        
        Returns:
            输出音频文件的绝对路径
        """
        try:
            from moviepy.editor import VideoFileClip
            
            # 加载视频
            video = VideoFileClip(video_path)
            
            # 提取音频
            audio = video.audio
            if audio is None:
                raise DocumentConvertError("该视频没有音频轨道")
            
            # 导出音频
            format_ext = options.get('format', 'mp3')
            export_kwargs = {'logger': None}
            if 'bitrate' in options:
                export_kwargs['bitrate'] = options['bitrate']
            
            audio.write_audiofile(output_path, **export_kwargs)
            
            video.close()
            audio.close()
            
            return os.path.abspath(output_path)
            
        except ImportError:
            raise DocumentConvertError("moviepy未安装。请运行: pip install moviepy")
        except Exception as e:
            raise DocumentConvertError(f"提取音频失败: {str(e)}")
    
    def merge_audio_video(self, video_path: str, audio_path: str, output_path: str, **options) -> str:
        """
        合并视频和音频
        
        Args:
            video_path: 视频文件路径
            audio_path: 音频文件路径
            output_path: 输出视频文件路径
            **options: 合并选项
        
        Returns:
            输出视频文件的绝对路径
        """
        try:
            from moviepy.editor import VideoFileClip, AudioFileClip
            
            # 加载视频和音频
            video = VideoFileClip(video_path)
            audio = AudioFileClip(audio_path)
            
            # 合并
            final_video = video.set_audio(audio)
            
            # 导出
            export_kwargs = {'logger': None}
            final_video.write_videofile(output_path, **export_kwargs)
            
            video.close()
            audio.close()
            final_video.close()
            
            return os.path.abspath(output_path)
            
        except ImportError:
            raise DocumentConvertError("moviepy未安装。请运行: pip install moviepy")
        except Exception as e:
            raise DocumentConvertError(f"合并音视频失败: {str(e)}")
    
    def get_audio_info(self, file_path: str) -> Dict[str, Any]:
        """获取音频文件信息"""
        try:
            from pydub import AudioSegment
            
            audio = AudioSegment.from_file(file_path)
            
            return {
                "type": "audio",
                "format": Path(file_path).suffix.lower().replace('.', ''),
                "duration_seconds": len(audio) / 1000,
                "duration_formatted": self._format_duration(len(audio) / 1000),
                "channels": audio.channels,
                "sample_rate": audio.frame_rate,
                "sample_width": audio.sample_width,
                "bitrate": audio.frame_rate * audio.channels * audio.sample_width * 8
            }
            
        except ImportError:
            raise DocumentReadError("pydub未安装。请运行: pip install pydub")
        except Exception as e:
            raise DocumentReadError(f"读取音频信息失败: {str(e)}")
    
    def transcribe_audio(self, audio_path: str, model_size: str = "base", 
                         language: str = "zh", verbose: bool = False) -> Dict[str, Any]:
        """
        使用Whisper转录音频（需要安装openai-whisper）
        
        Args:
            audio_path: 音频文件路径
            model_size: 模型大小 (tiny/base/small/medium/large)
            language: 语言代码 (zh/en/auto)
            verbose: 是否显示详细进度
            
        Returns:
            转录结果字典
        """
        try:
            import whisper
            import warnings
            warnings.filterwarnings('ignore')
            
            model = whisper.load_model(model_size)
            
            result = model.transcribe(
                audio_path,
                language=language if language != 'auto' else None,
                verbose=verbose
            )
            
            return result
            
        except ImportError:
            raise DocumentConvertError(
                "openai-whisper未安装，无法转录音频。请运行: pip install openai-whisper"
            )
        except Exception as e:
            raise DocumentConvertError(f"转录音频失败: {str(e)}")
    
    def transcribe_video(self, video_path: str, model_size: str = "base",
                         language: str = "zh") -> Dict[str, Any]:
        """
        转录视频中的音频（自动提取+转录）
        
        Args:
            video_path: 视频文件路径
            model_size: 模型大小
            language: 语言
            
        Returns:
            包含transcript/text/analysis的完整结果
        """
        import tempfile
        import os
        
        # 创建临时音频文件
        temp_audio = tempfile.NamedTemporaryFile(suffix='.wav', delete=False)
        temp_audio_path = temp_audio.name
        temp_audio.close()
        
        try:
            # 提取音频（使用ffmpeg，更可靠）
            import subprocess
            cmd = [
                'ffmpeg', '-i', video_path,
                '-vn', '-acodec', 'pcm_s16le',
                '-ar', '16000', '-ac', '1',
                temp_audio_path, '-y', '-loglevel', 'error'
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
            if result.returncode != 0:
                raise DocumentConvertError(f"提取音频失败: {result.stderr}")
            
            # 转录音频
            transcript = self.transcribe_audio(temp_audio_path, model_size, language)
            
            # 分析内容
            analysis = self._analyze_transcript(transcript['text'])
            
            return {
                'transcript': transcript,
                'text': transcript['text'],
                'analysis': analysis,
                'summary': self._generate_summary(transcript['text'])
            }
            
        finally:
            # 清理临时文件
            if os.path.exists(temp_audio_path):
                os.unlink(temp_audio_path)
    
    def _analyze_transcript(self, text: str) -> Dict[str, Any]:
        """分析转录文本内容"""
        import re
        
        char_count = len(text)
        
        # 检测语言
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        is_chinese = chinese_chars > char_count * 0.3
        
        # 内容分类
        categories = []
        category_keywords = {
            '会议/访谈': ['会议', '讨论', '汇报', '总结', '采访', '对话'],
            '教学/教程': ['教程', '讲解', '介绍', '如何', '方法', '步骤'],
            '新闻/资讯': ['新闻', '报道', '最新', '热点', '资讯'],
            '播客/对话': ['欢迎收听', '大家好', '我是', '本期', '嘉宾'],
            '职场/求职': ['求职', '面试', '工作', '公司', '职业', '简历'],
            '生活/旅行': ['旅行', '生活', '体验', '房车', '露营', '酒店']
        }
        
        for category, keywords in category_keywords.items():
            if any(kw in text for kw in keywords):
                categories.append(category)
        
        if not categories:
            categories.append('其他')
        
        # 提取关键句子
        sentences = re.split(r'[。！？\n]', text)
        key_sentences = [s.strip() for s in sentences if len(s.strip()) > 20][:5]
        
        return {
            'char_count': char_count,
            'language': '中文' if is_chinese else '英文',
            'categories': categories,
            'key_sentences': key_sentences
        }
    
    def _generate_summary(self, text: str, max_length: int = 500) -> str:
        """生成文本摘要"""
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        
        if not paragraphs:
            return text[:max_length]
        
        summary_parts = []
        current_length = 0
        
        for para in paragraphs:
            if current_length + len(para) > max_length:
                break
            summary_parts.append(para)
            current_length += len(para)
        
        return '\n'.join(summary_parts)
    
    def get_video_info(self, file_path: str) -> Dict[str, Any]:
        """获取视频文件信息"""
        try:
            from moviepy.editor import VideoFileClip
            
            video = VideoFileClip(file_path)
            
            info = {
                "type": "video",
                "format": Path(file_path).suffix.lower().replace('.', ''),
                "duration_seconds": video.duration,
                "duration_formatted": self._format_duration(video.duration),
                "fps": video.fps,
                "resolution": f"{int(video.w)}x{int(video.h)}",
                "width": int(video.w),
                "height": int(video.h),
                "has_audio": video.audio is not None,
                "size_bytes": os.path.getsize(file_path)
            }
            
            video.close()
            return info
            
        except ImportError:
            raise DocumentReadError("moviepy未安装。请运行: pip install moviepy")
        except Exception as e:
            raise DocumentReadError(f"读取视频信息失败: {str(e)}")
    
    def trim_audio(self, input_path: str, output_path: str, start_time: float, end_time: float) -> str:
        """
        裁剪音频
        
        Args:
            input_path: 输入音频文件路径
            output_path: 输出音频文件路径
            start_time: 开始时间（秒）
            end_time: 结束时间（秒）
        
        Returns:
            输出文件的绝对路径
        """
        try:
            from pydub import AudioSegment
            
            audio = AudioSegment.from_file(input_path)
            
            # 转换为毫秒
            start_ms = int(start_time * 1000)
            end_ms = int(end_time * 1000)
            
            # 裁剪
            trimmed = audio[start_ms:end_ms]
            trimmed.export(output_path)
            
            return os.path.abspath(output_path)
            
        except ImportError:
            raise DocumentConvertError("pydub未安装。请运行: pip install pydub")
        except Exception as e:
            raise DocumentConvertError(f"裁剪音频失败: {str(e)}")
    
    def trim_video(self, input_path: str, output_path: str, start_time: float, end_time: float) -> str:
        """
        裁剪视频
        
        Args:
            input_path: 输入视频文件路径
            output_path: 输出视频文件路径
            start_time: 开始时间（秒）
            end_time: 结束时间（秒）
        
        Returns:
            输出文件的绝对路径
        """
        try:
            from moviepy.editor import VideoFileClip
            
            video = VideoFileClip(input_path)
            trimmed = video.subclip(start_time, end_time)
            
            export_kwargs = {'logger': None}
            trimmed.write_videofile(output_path, **export_kwargs)
            
            video.close()
            trimmed.close()
            
            return os.path.abspath(output_path)
            
        except ImportError:
            raise DocumentConvertError("moviepy未安装。请运行: pip install moviepy")
        except Exception as e:
            raise DocumentConvertError(f"裁剪视频失败: {str(e)}")
    
    def _format_duration(self, seconds: float) -> str:
        """格式化时长为 HH:MM:SS"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        
        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"
        else:
            return f"{minutes:02d}:{secs:02d}"


class DocumentHub:
    """统一文档处理中心"""
    
    def __init__(self):
        self.supported_readers = {
            DocumentType.WORD: self._read_word,
            DocumentType.EXCEL: self._read_excel,
            DocumentType.PDF: self._read_pdf,
            DocumentType.MARKDOWN: self._read_markdown,
        }
        self.supported_writers = {
            DocumentType.WORD: self._write_word,
            DocumentType.EXCEL: self._write_excel,
            DocumentType.PDF: self._write_pdf,
            DocumentType.MARKDOWN: self._write_markdown,
        }
        # 音视频处理中心
        self.media = MediaHub()
    
    def convert_media(self, input_path: str, output_path: str, **options) -> str:
        """
        转换音视频格式
        
        Args:
            input_path: 输入文件路径
            output_path: 输出文件路径
            **options: 转换选项
        
        Returns:
            输出文件的绝对路径
        """
        input_type = get_document_type(input_path)
        output_type = get_document_type(output_path)
        
        # 音频 -> 音频
        if input_type == DocumentType.AUDIO and output_type == DocumentType.AUDIO:
            return self.media.convert_audio(input_path, output_path, **options)
        
        # 视频 -> 视频
        elif input_type == DocumentType.VIDEO and output_type == DocumentType.VIDEO:
            return self.media.convert_video(input_path, output_path, **options)
        
        # 视频 -> 音频（提取）
        elif input_type == DocumentType.VIDEO and output_type == DocumentType.AUDIO:
            return self.media.extract_audio_from_video(input_path, output_path, **options)
        
        else:
            raise DocumentConvertError(f"不支持的媒体转换: {input_type.value} -> {output_type.value}")
    
    def get_media_info(self, file_path: str) -> Dict[str, Any]:
        """获取音视频文件信息"""
        doc_type = get_document_type(file_path)
        
        if doc_type == DocumentType.AUDIO:
            return self.media.get_audio_info(file_path)
        elif doc_type == DocumentType.VIDEO:
            return self.media.get_video_info(file_path)
        else:
            raise DocumentReadError(f"不是音视频文件: {file_path}")
    
    def read(self, file_path: str, **options) -> Dict[str, Any]:
        """
        读取文档内容
        
        Args:
            file_path: 文件路径
            **options: 读取选项
                - extract_tables: 是否提取表格 (Excel/Word)
                - extract_images: 是否提取图片
                - max_pages: 最大读取页数 (PDF)
        
        Returns:
            包含文档内容的字典
        """
        if not os.path.exists(file_path):
            raise DocumentReadError(f"文件不存在: {sanitize_path(file_path)}")
        
        doc_type = get_document_type(file_path)
        if doc_type == DocumentType.UNKNOWN:
            raise DocumentReadError(f"不支持的文件类型: {Path(file_path).suffix}")
        
        if doc_type not in self.supported_readers:
            raise DocumentReadError(f"暂不支持读取该类型: {doc_type.value}")
        
        try:
            return self.supported_readers[doc_type](file_path, **options)
        except Exception as e:
            raise DocumentReadError(f"读取失败 [{doc_type.value}]: {str(e)}")
    
    def write(self, file_path: str, content: Any, **options) -> str:
        """
        写入文档
        
        Args:
            file_path: 输出文件路径
            content: 文档内容
            **options: 写入选项
        
        Returns:
            输出文件的绝对路径
        """
        # 确保目录存在
        os.makedirs(os.path.dirname(os.path.abspath(file_path)), exist_ok=True)
        
        doc_type = get_document_type(file_path)
        if doc_type == DocumentType.UNKNOWN:
            raise DocumentWriteError(f"不支持的文件类型: {Path(file_path).suffix}")
        
        if doc_type not in self.supported_writers:
            raise DocumentWriteError(f"暂不支持写入该类型: {doc_type.value}")
        
        try:
            return self.supported_writers[doc_type](file_path, content, **options)
        except Exception as e:
            raise DocumentWriteError(f"写入失败 [{doc_type.value}]: {str(e)}")
    
    def convert(self, input_path: str, output_path: str, **options) -> str:
        """
        转换文档格式
        
        Args:
            input_path: 输入文件路径
            output_path: 输出文件路径
            **options: 转换选项
        
        Returns:
            输出文件的绝对路径
        """
        input_type = get_document_type(input_path)
        output_type = get_document_type(output_path)
        
        # Word -> PDF
        if input_type == DocumentType.WORD and output_type == DocumentType.PDF:
            return self._word_to_pdf(input_path, output_path, **options)
        
        # Markdown -> Word
        elif input_type == DocumentType.MARKDOWN and output_type == DocumentType.WORD:
            return self._markdown_to_word(input_path, output_path, **options)
        
        else:
            raise DocumentConvertError(f"不支持的转换: {input_type.value} -> {output_type.value}")
    
    def get_info(self, file_path: str) -> DocumentInfo:
        """获取文档信息"""
        if not os.path.exists(file_path):
            raise DocumentReadError(f"文件不存在: {sanitize_path(file_path)}")
        
        doc_type = get_document_type(file_path)
        size = os.path.getsize(file_path)
        
        info = DocumentInfo(
            path=file_path,
            doc_type=doc_type,
            size=size
        )
        
        # 根据类型获取更多信息
        if doc_type == DocumentType.WORD:
            info = self._get_word_info(file_path, info)
        elif doc_type == DocumentType.PDF:
            info = self._get_pdf_info(file_path, info)
        
        return info
    
    # ========== Word处理 ==========
    
    def _read_word(self, file_path: str, **options) -> Dict[str, Any]:
        """读取Word文档"""
        from docx import Document
        
        doc = Document(file_path)
        result = {
            "type": "word",
            "paragraphs": [],
            "tables": [],
            "metadata": {}
        }
        
        # 读取段落
        for para in doc.paragraphs:
            if para.text.strip():
                result["paragraphs"].append(para.text)
        
        # 读取表格（如果启用）
        if options.get("extract_tables", True):
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                result["tables"].append(table_data)
        
        result["metadata"]["paragraph_count"] = len(doc.paragraphs)
        result["metadata"]["table_count"] = len(doc.tables)
        
        return result
    
    def _write_word(self, file_path: str, content: Dict[str, Any], **options) -> str:
        """写入Word文档"""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        
        doc = Document()
        
        # 设置默认字体
        doc.styles['Normal'].font.name = options.get('font_name', 'SimSun')
        doc.styles['Normal']._element.rPr.rFonts.set(
            qn('w:eastAsia'), 
            options.get('font_name', 'SimSun')
        )
        
        # 写入标题
        if "title" in content:
            title = doc.add_heading(content["title"], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 写入段落
        for para in content.get("paragraphs", []):
            if isinstance(para, dict):
                p = doc.add_paragraph(para.get("text", ""))
                if para.get("bold"):
                    p.runs[0].bold = True
                if para.get("style"):
                    p.style = para["style"]
            else:
                doc.add_paragraph(str(para))
        
        # 写入表格
        for table_data in content.get("tables", []):
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                for i, row_data in enumerate(table_data):
                    for j, cell_text in enumerate(row_data):
                        table.rows[i].cells[j].text = str(cell_text)
        
        doc.save(file_path)
        return os.path.abspath(file_path)
    
    def _get_word_info(self, file_path: str, info: DocumentInfo) -> DocumentInfo:
        """获取Word文档详细信息"""
        from docx import Document
        
        doc = Document(file_path)
        info.paragraphs = len(doc.paragraphs)
        info.tables = len(doc.tables)
        
        return info
    
    # ========== Excel处理 ==========
    
    def _read_excel(self, file_path: str, **options) -> Dict[str, Any]:
        """读取Excel文档"""
        import pandas as pd
        
        # 读取所有sheet
        xls = pd.ExcelFile(file_path)
        result = {
            "type": "excel",
            "sheets": {},
            "metadata": {
                "sheet_names": xls.sheet_names,
                "sheet_count": len(xls.sheet_names)
            }
        }
        
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            result["sheets"][sheet_name] = {
                "data": df.to_dict(orient='records'),
                "headers": df.columns.tolist(),
                "row_count": len(df),
                "column_count": len(df.columns)
            }
        
        return result
    
    def _write_excel(self, file_path: str, content: Dict[str, Any], **options) -> str:
        """写入Excel文档"""
        import pandas as pd
        
        with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
            for sheet_name, sheet_data in content.get("sheets", {}).items():
                df = pd.DataFrame(sheet_data.get("data", []))
                df.to_excel(writer, sheet_name=sheet_name, index=False)
        
        return os.path.abspath(file_path)
    
    # ========== PDF处理 ==========
    
    def _read_pdf(self, file_path: str, **options) -> Dict[str, Any]:
        """读取PDF文档"""
        import pdfplumber
        
        result = {
            "type": "pdf",
            "pages": [],
            "metadata": {}
        }
        
        with pdfplumber.open(file_path) as pdf:
            max_pages = options.get("max_pages", len(pdf.pages))
            
            for i, page in enumerate(pdf.pages[:max_pages]):
                page_text = page.extract_text()
                page_tables = page.extract_tables() if options.get("extract_tables", True) else []
                
                result["pages"].append({
                    "page_num": i + 1,
                    "text": page_text or "",
                    "tables": page_tables
                })
            
            result["metadata"]["total_pages"] = len(pdf.pages)
            result["metadata"]["read_pages"] = min(max_pages, len(pdf.pages))
        
        return result
    
    def _write_pdf(self, file_path: str, content: Any, **options) -> str:
        """写入PDF文档（简单文本）"""
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        # 注册中文字体（如果需要）
        try:
            pdfmetrics.registerFont(TTFont('SimSun', '/System/Library/Fonts/STHeiti Light.ttc'))
        except:
            pass
        
        c = canvas.Canvas(file_path, pagesize=A4)
        width, height = A4
        
        y = height - 50
        for text in content.get("texts", []):
            c.drawString(50, y, str(text))
            y -= 20
            if y < 50:
                c.showPage()
                y = height - 50
        
        c.save()
        return os.path.abspath(file_path)
    
    def _get_pdf_info(self, file_path: str, info: DocumentInfo) -> DocumentInfo:
        """获取PDF文档详细信息"""
        import pdfplumber
        
        with pdfplumber.open(file_path) as pdf:
            info.pages = len(pdf.pages)
        
        return info
    
    # ========== Markdown处理 ==========
    
    def _read_markdown(self, file_path: str, **options) -> Dict[str, Any]:
        """读取Markdown文档"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return {
            "type": "markdown",
            "content": content,
            "metadata": {
                "line_count": len(content.split('\n')),
                "char_count": len(content)
            }
        }
    
    def _write_markdown(self, file_path: str, content: str, **options) -> str:
        """写入Markdown文档"""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return os.path.abspath(file_path)
    
    # ========== 格式转换 ==========
    
    def _word_to_pdf(self, input_path: str, output_path: str, **options) -> str:
        """Word转PDF（使用LibreOffice）"""
        import subprocess
        
        output_dir = os.path.dirname(os.path.abspath(output_path))
        
        try:
            # 使用LibreOffice命令行转换
            cmd = [
                'soffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', output_dir,
                input_path
            ]
            subprocess.run(cmd, check=True, capture_output=True, timeout=60)
            
            # 重命名输出文件
            base_name = Path(input_path).stem
            generated_pdf = os.path.join(output_dir, f"{base_name}.pdf")
            
            if generated_pdf != output_path and os.path.exists(generated_pdf):
                os.rename(generated_pdf, output_path)
            
            return os.path.abspath(output_path)
            
        except subprocess.CalledProcessError as e:
            raise DocumentConvertError(f"LibreOffice转换失败: {e.stderr.decode()}")
        except FileNotFoundError:
            raise DocumentConvertError("LibreOffice未安装，无法转换Word到PDF")
    
    def _markdown_to_word(self, input_path: str, output_path: str, **options) -> str:
        """Markdown转Word"""
        # 先读取Markdown
        md_content = self._read_markdown(input_path)
        
        # 简单解析（可以扩展使用markdown库）
        lines = md_content["content"].split('\n')
        word_content = {
            "paragraphs": [],
            "tables": []
        }
        
        for line in lines:
            line = line.strip()
            if line.startswith('# '):
                word_content["paragraphs"].append({"text": line[2:], "style": "Heading 1"})
            elif line.startswith('## '):
                word_content["paragraphs"].append({"text": line[3:], "style": "Heading 2"})
            elif line.startswith('### '):
                word_content["paragraphs"].append({"text": line[4:], "style": "Heading 3"})
            elif line.startswith('- '):
                word_content["paragraphs"].append({"text": "• " + line[2:]})
            elif line:
                word_content["paragraphs"].append(line)
        
        return self._write_word(output_path, word_content, **options)

    # ========== 飞书协作接口 ==========
    
    def _check_feishu_skills(self) -> bool:
        """检查飞书技能是否可用"""
        try:
            self._load_feishu_creator()
            self._load_feishu_converter()
            return True
        except ImportError:
            return False
    
    def _load_feishu_creator(self):
        """加载 feishu-doc-creator（处理连字符目录名）"""
        import importlib.util
        from pathlib import Path
        
        current_file = Path(__file__).resolve()
        creator_path = current_file.parent.parent / 'feishu-doc-creator' / '__init__.py'
        
        spec = importlib.util.spec_from_file_location('feishu_doc_creator', str(creator_path))
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        
        # 注册到 sys.modules
        import sys
        sys.modules['feishu_doc_creator'] = module
        return module
    
    def _load_feishu_converter(self):
        """加载 feishu-doc-converter（处理连字符目录名）"""
        import importlib.util
        from pathlib import Path
        
        current_file = Path(__file__).resolve()
        converter_path = current_file.parent.parent / 'feishu-doc-converter' / '__init__.py'
        
        spec = importlib.util.spec_from_file_location('feishu_doc_converter', str(converter_path))
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        
        # 注册到 sys.modules
        import sys
        sys.modules['feishu_doc_converter'] = module
        return module
    
    def _ensure_feishu_imports(self):
        """确保飞书技能已导入"""
        import sys
        if 'feishu_doc_creator' not in sys.modules:
            self._load_feishu_creator()
        if 'feishu_doc_converter' not in sys.modules:
            self._load_feishu_converter()
    
    def to_feishu(self, file_path: str, title: str = None, 
                  doc_type: str = "drive", **options) -> dict:
        """
        将本地文档上传到飞书
        
        Args:
            file_path: 本地文件路径（Word/Excel/PDF/Markdown）
            title: 飞书文档标题（默认为文件名）
            doc_type: "drive"(云盘) 或 "wiki"(知识库)
            **options: 
                - folder_token: 云盘文件夹token（drive类型）
                - parent_node_token: 知识库父节点token（wiki类型）
        
        Returns:
            飞书文档信息字典，包含 url、document_id 等
        
        Raises:
            DocumentError: 飞书技能未安装或上传失败
        
        Example:
            >>> hub = get_hub()
            >>> result = hub.to_feishu("report.docx", title="月度报告")
            >>> print(result['url'])
        """
        if not self._check_feishu_skills():
            raise DocumentError(
                "飞书技能未安装，无法上传到飞书。\n"
                "请确保 skills/feishu-doc-creator 和 skills/feishu-doc-converter 存在。"
            )
        
        self._ensure_feishu_imports()
        
        from feishu_doc_creator import create_drive_doc, create_wiki_doc
        
        if not os.path.exists(file_path):
            raise DocumentReadError(f"文件不存在: {sanitize_path(file_path)}")
        
        # 确定标题
        title = title or Path(file_path).stem
        
        # 根据文件类型处理
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.md':
            # Markdown 直接读取
            with open(file_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
        else:
            # 其他格式：转换为 Markdown
            try:
                # 创建临时 MD 文件
                md_temp = file_path + '.tmp.md'
                self.convert(file_path, md_temp)
                with open(md_temp, 'r', encoding='utf-8') as f:
                    md_content = f.read()
                os.remove(md_temp)
            except DocumentConvertError as e:
                # 如果转换失败，尝试直接读取文本
                doc_content = self.read(file_path)
                md_content = self._doc_to_markdown(doc_content, title)
        
        # 上传到飞书
        try:
            if doc_type == "drive":
                result = create_drive_doc(
                    title=title,
                    content=md_content,
                    folder_token=options.get('folder_token')
                )
            else:  # wiki
                result = create_wiki_doc(
                    title=title,
                    content=md_content,
                    parent_node_token=options.get('parent_node_token')
                )
            return result
        except Exception as e:
            raise DocumentError(f"上传到飞书失败: {str(e)}")
    
    def from_feishu(self, doc_url: str, output_path: str = None) -> str:
        """
        从飞书文档下载到本地
        
        Args:
            doc_url: 飞书文档URL（支持 docx/xxx 和 wiki/xxx 格式）
            output_path: 输出文件路径（可选，默认返回 Markdown 内容）
                     支持 .md、.docx、.pdf 等格式
        
        Returns:
            如果 output_path 为 None，返回 Markdown 字符串
            否则返回输出文件的绝对路径
        
        Raises:
            DocumentError: 飞书技能未安装或下载失败
        
        Example:
            >>> hub = get_hub()
            >>> # 获取 Markdown 内容
            >>> md_content = hub.from_feishu("https://feishu.cn/docx/xxx")
            >>> # 保存为 Word
            >>> hub.from_feishu("https://feishu.cn/docx/xxx", "output.docx")
        """
        if not self._check_feishu_skills():
            raise DocumentError(
                "飞书技能未安装，无法从飞书下载。\n"
                "请确保 skills/feishu-doc-creator 和 skills/feishu-doc-converter 存在。"
            )
        
        self._ensure_feishu_imports()
        
        from feishu_doc_converter import doc_to_md, url_to_md
        
        from skills.feishu_doc_converter import doc_to_md, url_to_md
        
        # 根据 URL 类型选择导出方式
        if 'wiki' in doc_url:
            # 知识库文档使用 url_to_md（browser 方式）
            md_content = url_to_md(doc_url)
        else:
            # 云盘文档使用 doc_to_md（API 方式）
            md_content = doc_to_md(doc_url)
        
        # 如果没有指定输出路径，直接返回 Markdown 内容
        if output_path is None:
            return md_content
        
        # 根据输出格式处理
        output_ext = Path(output_path).suffix.lower()
        
        if output_ext == '.md':
            # 直接保存为 Markdown
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            return os.path.abspath(output_path)
        
        # 其他格式：先保存为临时 MD，再转换
        md_temp = output_path + '.tmp.md'
        try:
            with open(md_temp, 'w', encoding='utf-8') as f:
                f.write(md_content)
            
            # 转换为目标格式
            self.convert(md_temp, output_path)
            return os.path.abspath(output_path)
        finally:
            # 清理临时文件
            if os.path.exists(md_temp):
                os.remove(md_temp)
    
    def _doc_to_markdown(self, doc_content: Dict[str, Any], title: str = None) -> str:
        """将文档内容字典转换为 Markdown 字符串"""
        lines = []
        
        # 添加标题
        if title:
            lines.append(f"# {title}")
            lines.append("")
        
        # 处理段落
        if "paragraphs" in doc_content:
            for para in doc_content["paragraphs"]:
                if isinstance(para, dict):
                    text = para.get("text", "")
                    style = para.get("style", "")
                    if style.startswith("Heading"):
                        level = style.split()[-1]
                        lines.append(f"{'#' * int(level)} {text}")
                    elif para.get("bold"):
                        lines.append(f"**{text}**")
                    else:
                        lines.append(text)
                elif isinstance(para, str):
                    lines.append(para)
                lines.append("")
        
        # 处理页面（PDF）
        if "pages" in doc_content:
            for page in doc_content["pages"]:
                if page.get("text"):
                    lines.append(page["text"])
                    lines.append("")
        
        # 处理表格
        if "tables" in doc_content and doc_content["tables"]:
            for table in doc_content["tables"]:
                if table:
                    lines.append("")
                    # 简单的表格 Markdown 转换
                    for i, row in enumerate(table):
                        if isinstance(row, list):
                            row_text = " | ".join(str(cell) for cell in row)
                            lines.append(f"| {row_text} |")
                            if i == 0:  # 表头分隔符
                                lines.append("| " + " | ".join(["---"] * len(row)) + " |")
                    lines.append("")
        
        return "\n".join(lines)

# 全局实例
_hub = None

def get_hub() -> DocumentHub:
    """获取DocumentHub单例"""
    global _hub
    if _hub is None:
        _hub = DocumentHub()
    return _hub

# 便捷函数
def read(file_path: str, **options) -> Dict[str, Any]:
    """读取文档"""
    return get_hub().read(file_path, **options)

def write(file_path: str, content: Any, **options) -> str:
    """写入文档"""
    return get_hub().write(file_path, content, **options)

def convert(input_path: str, output_path: str, **options) -> str:
    """转换文档"""
    return get_hub().convert(input_path, output_path, **options)

def get_info(file_path: str) -> DocumentInfo:
    """获取文档信息"""
    return get_hub().get_info(file_path)

def to_feishu(file_path: str, title: str = None, doc_type: str = "drive", **options) -> dict:
    """上传本地文档到飞书"""
    return get_hub().to_feishu(file_path, title, doc_type, **options)

def from_feishu(doc_url: str, output_path: str = None) -> str:
    """从飞书文档下载到本地"""
    return get_hub().from_feishu(doc_url, output_path)

if __name__ == "__main__":
    # 测试代码
    hub = DocumentHub()
    print("Document Hub initialized successfully")
    print(f"Supported readers: {list(hub.supported_readers.keys())}")
    print(f"Supported writers: {list(hub.supported_writers.keys())}")
    print(f"Feishu skills available: {hub._check_feishu_skills()}")
